"""Typed, fail-closed document extraction for Emilia and the request broker."""

from __future__ import annotations

import csv
import json
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Optional


TEXT_EXTENSIONS = {
    ".txt", ".md", ".json", ".log", ".py", ".js", ".ts", ".tsx",
    ".jsx", ".html", ".css", ".yaml", ".yml", ".toml", ".xml",
}
SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | {
    ".pdf", ".docx", ".xlsx", ".xlsm", ".csv", ".tsv",
}


@dataclass(frozen=True)
class ExtractionResult:
    ok: bool
    supported: bool
    text: str = ""
    extension: str = ""
    pages: Optional[int] = None
    truncated: bool = False
    warnings: list[str] = field(default_factory=list)
    error_code: Optional[str] = None
    error_message: Optional[str] = None

    def public_error(self) -> str:
        if not self.supported:
            return "Формат файла пока не поддерживается для текстового анализа."
        if self.error_code == "empty_document":
            return "В документе не найден текст. Возможно, это скан; для него нужен OCR."
        return "Не удалось безопасно извлечь текст из документа. Файл сохранён, можно повторить обработку."

    def to_dict(self) -> dict:
        return asdict(self)


def _limit(text: str, max_chars: int) -> tuple[str, bool]:
    if len(text) <= max_chars:
        return text, False
    return text[:max_chars], True


def _extract_pdf(path: Path) -> tuple[str, int, list[str]]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    sections = []
    warnings = []
    for index, page in enumerate(reader.pages, start=1):
        page_text = (page.extract_text() or "").strip()
        if page_text:
            sections.append(f"[Страница {index}]\n{page_text}")
    if not sections and reader.pages:
        warnings.append("requires_ocr")
    return "\n\n".join(sections), len(reader.pages), warnings


def _extract_docx(path: Path) -> str:
    import docx

    document = docx.Document(str(path))
    lines = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                lines.append("\t".join(values))
    return "\n".join(lines)


def _extract_workbook(path: Path) -> str:
    import openpyxl

    workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
    try:
        lines = []
        for sheet in workbook.worksheets:
            lines.append(f"[Лист: {sheet.title}]")
            for row in sheet.iter_rows(values_only=True):
                values = [str(value) for value in row if value is not None]
                if values:
                    lines.append("\t".join(values))
        return "\n".join(lines)
    finally:
        workbook.close()


def _extract_delimited(path: Path, delimiter: str) -> str:
    lines = []
    with path.open("r", encoding="utf-8-sig", errors="replace", newline="") as handle:
        for row in csv.reader(handle, delimiter=delimiter):
            if any(cell.strip() for cell in row):
                lines.append("\t".join(row))
    return "\n".join(lines)


def extract_document(path: str | Path, max_chars: int = 120_000) -> ExtractionResult:
    """Extract document text without ever returning an exception as document content."""
    file_path = Path(path)
    extension = file_path.suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        return ExtractionResult(ok=False, supported=False, extension=extension, error_code="unsupported")

    try:
        pages = None
        warnings: list[str] = []
        if extension == ".pdf":
            text, pages, warnings = _extract_pdf(file_path)
        elif extension == ".docx":
            text = _extract_docx(file_path)
        elif extension in {".xlsx", ".xlsm"}:
            text = _extract_workbook(file_path)
        elif extension == ".csv":
            text = _extract_delimited(file_path, ",")
        elif extension == ".tsv":
            text = _extract_delimited(file_path, "\t")
        else:
            text = file_path.read_text(encoding="utf-8", errors="replace")
    except Exception as error:
        # Extraction libraries expose several format-specific exception trees.
        # This boundary must fail closed so no parser error becomes model input.
        return ExtractionResult(
            ok=False,
            supported=True,
            extension=extension,
            error_code="extraction_failed",
            error_message=f"{error.__class__.__name__}: {error}",
        )

    text = text.strip()
    if not text:
        return ExtractionResult(
            ok=False,
            supported=True,
            extension=extension,
            pages=pages,
            warnings=warnings,
            error_code="empty_document",
        )
    text, truncated = _limit(text, max_chars)
    if truncated:
        warnings.append("truncated")
    return ExtractionResult(
        ok=True,
        supported=True,
        text=text,
        extension=extension,
        pages=pages,
        truncated=truncated,
        warnings=warnings,
    )
