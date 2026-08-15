from datetime import datetime, timedelta, timezone
from pathlib import Path

import artifact_store
import document_pipeline
import orchestrator


def test_text_extraction_returns_typed_success(tmp_path):
    path = tmp_path / "brief.txt"
    path.write_text("Главный риск: срок поставки.", encoding="utf-8")

    result = document_pipeline.extract_document(path)

    assert result.ok is True
    assert result.supported is True
    assert "срок поставки" in result.text
    assert result.error_message is None


def test_unsupported_document_is_not_model_content(tmp_path):
    path = tmp_path / "archive.bin"
    path.write_bytes(b"not a document")

    result = document_pipeline.extract_document(path)

    assert result.ok is False
    assert result.supported is False
    assert result.text == ""
    assert "не поддерживается" in result.public_error()


def test_extraction_error_is_separate_from_document_text(monkeypatch, tmp_path):
    path = tmp_path / "broken.pdf"
    path.write_bytes(b"broken")
    monkeypatch.setattr(
        document_pipeline,
        "_extract_pdf",
        lambda _path: (_ for _ in ()).throw(ImportError("No module named pypdf")),
    )

    result = document_pipeline.extract_document(path)

    assert result.ok is False
    assert result.text == ""
    assert result.error_code == "extraction_failed"
    assert "pypdf" in result.error_message
    assert "pypdf" not in result.public_error()


def test_docx_extracts_paragraphs_and_tables(tmp_path):
    import docx

    path = tmp_path / "contract.docx"
    document = docx.Document()
    document.add_paragraph("Договор поставки")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Срок"
    table.cell(0, 1).text = "30 дней"
    document.save(path)

    result = document_pipeline.extract_document(path)

    assert result.ok is True
    assert "Договор поставки" in result.text
    assert "Срок\t30 дней" in result.text


def test_artifact_store_preserves_active_document(monkeypatch, tmp_path):
    monkeypatch.setenv("AMORI_ARTIFACT_ROOT", str(tmp_path / "artifacts"))
    source = tmp_path / "source.txt"
    source.write_text("private content", encoding="utf-8")

    artifact = artifact_store.store_file(source, "../contract.txt", "user-1", source="telegram")
    artifact = artifact_store.attach_extracted_text(artifact, "extracted content")
    active = artifact_store.get_active("user-1")

    assert active is not None
    assert active.id == artifact.id
    assert Path(active.stored_path).name == "contract.txt"
    assert Path(active.extracted_text_path).read_text(encoding="utf-8") == "extracted content"
    assert Path(active.stored_path).stat().st_mode & 0o077 == 0


def test_artifact_cleanup_removes_expired_data(monkeypatch, tmp_path):
    monkeypatch.setenv("AMORI_ARTIFACT_ROOT", str(tmp_path / "artifacts"))
    source = tmp_path / "source.txt"
    source.write_text("content", encoding="utf-8")
    artifact = artifact_store.store_file(
        source, "source.txt", "user-1", source="telegram", retention_days=1
    )

    removed = artifact_store.cleanup_expired(datetime.now(timezone.utc) + timedelta(days=2))

    assert removed == 1
    assert artifact_store.get_artifact(artifact.id) is None


def test_document_analysis_uses_subscription_router_first(monkeypatch):
    extraction = document_pipeline.ExtractionResult(
        ok=True, supported=True, text="[Страница 1]\nСрок договора 30 дней", extension=".pdf", pages=1
    )
    monkeypatch.setattr(orchestrator.llm, "smart_router_answer", lambda *_a, **_k: "Риск: короткий срок")
    monkeypatch.setattr(
        orchestrator.llm,
        "qwen_answer",
        lambda *_a, **_k: (_ for _ in ()).throw(AssertionError("fallback must not run")),
    )

    result = orchestrator.analyze_document(extraction, "contract.pdf", "найди риски")

    assert result == "Риск: короткий срок"


def test_active_document_is_added_only_for_file_followup(monkeypatch, tmp_path):
    monkeypatch.setenv("AMORI_ARTIFACT_ROOT", str(tmp_path / "artifacts"))
    source = tmp_path / "contract.txt"
    source.write_text("Срок оплаты 10 дней", encoding="utf-8")
    artifact = artifact_store.store_file(source, "contract.txt", "user-1", source="telegram")
    artifact_store.attach_extracted_text(artifact, "Срок оплаты 10 дней")

    assert "Срок оплаты" in orchestrator._active_artifact_context("user-1", "Что по этому договору?")
    assert orchestrator._active_artifact_context("user-1", "Какая сегодня погода?") == ""
